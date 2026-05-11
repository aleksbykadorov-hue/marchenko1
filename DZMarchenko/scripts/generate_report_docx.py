# -*- coding: utf-8 -*-
"""Генерация отчёта по курсовому проекту (Word .docx). Запуск: python scripts/generate_report_docx.py"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
# Два имени: кириллица для сдачи и латиница — если проводник/архивы коверкают имя
OUT_RU = ROOT / "Отчет_Чат_приложение.docx"
OUT_EN = ROOT / "Otchet_Chat_prilozhenie.docx"


def shot(p: Document, text: str) -> None:
    """Маркер места для вставки скриншота."""
    r = p.add_paragraph()
    r.paragraph_format.left_indent = Pt(12)
    run = r.add_run(f"【СКРИНШОТ】 {text}")
    run.italic = True
    run.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def main() -> None:
    doc = Document()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "ОТЧЁТ\nпо проекту\n«Чат-приложение: система обмена сообщениями с возможностью создания групповых чатов»"
    )
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Дисциплина: современные технологии разработки ПО (пример)")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Выполнил: _________________________ (ФИО, группа)")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("Проверил: _________________________")

    doc.add_page_break()

    # --- Введение ---
    add_heading(doc, "Введение", 1)
    doc.add_paragraph(
        "В отчёте представлены результаты проектирования и реализации распределённого "
        "чат-приложения с групповыми чатами, аутентификацией по JWT и микросервисной архитектурой. "
        "Материал структурирован по этапам учебного задания: проектирование, реализация API и "
        "контейнеризация, микросервисы и интеграция, оптимизация и развёртывание."
    )

    # ========== ЭТАП 1: Проектирование ==========
    add_heading(doc, "1. Проектирование архитектуры", 1)

    add_heading(doc, "1.1. Выбор архитектуры: монолит или микросервисы", 2)
    doc.add_paragraph(
        "Для системы выбрана микросервисная архитектура (минимальный набор сервисов) с "
        "синхронным взаимодействием по REST API.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Обоснование: разделение ответственности между доменом пользователей/аутентификации и "
        "доменом чатов; возможность независимого масштабирования и развёртывания; наглядная "
        "интеграция для учебного проекта (OpenAPI, HTTP).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Монолитная альтернатива могла бы упростить развёртывание на ранней стадии, но не "
        "соответствует целям этапа по микросервисам и разделению bounded context.",
        style="List Bullet",
    )

    add_heading(doc, "1.2. Основные компоненты и взаимодействие", 2)
    doc.add_paragraph("Компоненты системы:", style="List Number")
    doc.add_paragraph(
        "user-service — регистрация, вход, выпуск JWT, профиль пользователя; собственная БД PostgreSQL (user_db).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "chat-service — групповые чаты, участники, сообщения (CRUD); БД PostgreSQL (chat_db); "
        "валидация JWT с тем же секретом, что и у user-service.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "notification-service — приём запросов на уведомление (учебная заглушка), сохранение "
        "события в БД; вызывается из chat-service при создании сообщения (best-effort).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Клиент — веб-интерфейс (статические файлы chat-service) и/или прямые запросы к OpenAPI.",
        style="List Bullet",
    )

    add_heading(doc, "1.3. Архитектурная диаграмма", 2)
    doc.add_paragraph(
        "Ожидаемый результат задания — архитектурная диаграмма. В репозитории она задана в формате "
        "Mermaid (файл docs/architecture.md); ниже — текстовый эквивалент для отчёта."
    )
    doc.add_paragraph(
        "Клиент (браузер) → REST → user-service и chat-service.\n"
        "chat-service → REST → notification-service.\n"
        "user-service → PostgreSQL (user_db); chat-service → PostgreSQL (chat_db); "
        "notification-service → PostgreSQL (notif_db)."
    )
    shot(
        doc,
        "Экспорт диаграммы из docs/architecture.md (Mermaid) в PNG/SVG и вставка сюда; "
        "либо скриншот из редактора диаграмм (draw.io) по той же схеме.",
    )

    add_heading(doc, "1.4. Технологии и фреймворки", 2)
    doc.add_paragraph("API: FastAPI (асинхронная основа Starlette, автоматическая OpenAPI/Swagger).", style="List Bullet")
    doc.add_paragraph("БД: PostgreSQL 16 в Docker Compose; ORM — SQLAlchemy 2.x.", style="List Bullet")
    doc.add_paragraph("Валидация и схемы: Pydantic.", style="List Bullet")
    doc.add_paragraph("Аутентификация: JWT (алгоритм HS256), пароли — хэширование (bcrypt).", style="List Bullet")
    doc.add_paragraph("Тесты: pytest, TestClient (httpx).", style="List Bullet")
    doc.add_paragraph("Контейнеризация: Docker, Docker Compose.", style="List Bullet")
    doc.add_paragraph(
        "Альтернативы из формулировки задания: Flask вместо FastAPI — возможны, но для OpenAPI и "
        "типизации выбран FastAPI; MongoDB вместо PostgreSQL — не использовалась из-за реляционной "
        "модели чатов и участников.",
        style="List Bullet",
    )

    add_heading(doc, "1.5. Архитектурные паттерны", 2)
    doc.add_paragraph(
        "Слоистая организация API (аналог MVC для веб-API): маршруты в main.py (контроллеры), "
        "модели SQLAlchemy и схемы Pydantic (представление данных), доступ к данным через сессию БД.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Dependency Injection / Factory: фабрика сессий get_db() и зависимости FastAPI для внедрения БД и user_id.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Singleton (ограниченно): объект настроек Settings как единый источник конфигурации процесса.",
        style="List Bullet",
    )

    add_heading(doc, "1.6. Техническое задание (функционал и требования)", 2)
    doc.add_paragraph(
        "Полное ТЗ зафиксировано в репозитории в файле docs/technical_spec.md. Краткое содержание:"
    )
    doc.add_paragraph("Назначение: обмен сообщениями, групповые чаты, аутентификация.", style="List Bullet")
    doc.add_paragraph(
        "user-service: регистрация (email, пароль), логин → JWT, GET /me по токену.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "chat-service: создание чата, список чатов пользователя, добавление участника (владелец), "
        "отправка и просмотр сообщений с пагинацией, удаление сообщения автором.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "notification-service: POST /notify — приём уведомления (демо), сохранение в БД.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Нефункциональные: хэш паролей, секреты через переменные окружения, валидация входа, OpenAPI, тесты.",
        style="List Bullet",
    )

    doc.add_page_break()

    # ========== ЭТАП 3 ==========
    add_heading(doc, "2. Этап 3. Реализация основного API и контейнеризация", 1)

    add_heading(doc, "2.1. Реализация API (CRUD, JWT)", 2)
    doc.add_paragraph(
        "Реализованы эндпоинты для сущностей чата и пользователя в соответствии с ТЗ. "
        "Операции создания, чтения и удаления покрывают основной сценарий; отдельные поля "
        "редактируются через соответствующие операции (например, удаление сообщения)."
    )
    doc.add_paragraph("user-service:", style="List Bullet")
    doc.add_paragraph("POST /auth/register, POST /auth/login, GET /me", style="List Bullet 2")
    doc.add_paragraph("chat-service:", style="List Bullet")
    doc.add_paragraph(
        "POST /chats, GET /chats, POST /chats/{id}/members, POST/GET /chats/{id}/messages, DELETE /messages/{id}",
        style="List Bullet 2",
    )
    doc.add_paragraph("notification-service: POST /notify", style="List Bullet")
    shot(doc, "Swagger user-service: страница http://localhost:8001/docs (общий вид списка эндпоинтов).")
    shot(doc, "Swagger chat-service: http://localhost:8002/docs — эндпоинты чатов и сообщений.")
    shot(doc, "Swagger notification-service: http://localhost:8003/docs — POST /notify.")

    add_heading(doc, "2.2. Тестирование API (pytest)", 2)
    doc.add_paragraph(
        "Написаны тесты на pytest с использованием TestClient: проверка регистрации, логина и /me; "
        "сценарий CRUD чата (создание, список, участник, сообщение, список сообщений, удаление); "
        "уведомления. Файлы: services/*/tests/."
    )
    shot(
        doc,
        "Терминал: вывод команды pytest из папки user_service (или общий скрин успешного прогона по всем сервисам).",
    )
    shot(doc, "Терминал: pytest для chat_service и/или notification_service (зелёные тесты).")

    add_heading(doc, "2.3. Контейнеризация", 2)
    doc.add_paragraph(
        "Для каждого сервиса создан Dockerfile. В корне проекта — docker-compose.yml: три экземпляра "
        "PostgreSQL и три приложения; настроены переменные окружения, порты 8001–8003, healthcheck БД "
        "и зависимость сервисов от готовности PostgreSQL."
    )
    shot(doc, "Docker Desktop: список контейнеров проекта dzmarchenko — все running.")
    shot(doc, "Терминал: фрагмент лога docker compose up — строки Uvicorn running / Application startup complete.")

    doc.add_page_break()

    # ========== ЭТАП 4 ==========
    add_heading(doc, "3. Этап 4. Разработка и интеграция микросервисов", 1)

    add_heading(doc, "3.1. Разработка микросервисов", 2)
    doc.add_paragraph(
        "В задании перечислены примеры микросервисов: пользователи, платежи, уведомления. "
        "В рамках проекта реализованы:"
    )
    doc.add_paragraph(
        "Микросервис пользователей — регистрация, аутентификация, JWT (user-service).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Микросервис уведомлений — приём событий по REST, сохранение в БД (notification-service).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Микросервис платежей не реализовывался: не входит в предметную область учебного чат-приложения; "
        "при необходимости может быть добавлен как отдельный сервис с API оплаты и webhooks.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Основной домен чатов выделен в chat-service (групповые чаты и сообщения).",
        style="List Bullet",
    )

    add_heading(doc, "3.2. Интеграция микросервисов", 2)
    doc.add_paragraph(
        "Клиент обращается к user-service и chat-service по REST. chat-service при создании сообщения "
        "выполняет HTTP POST на notification-service (URL из переменной NOTIFICATION_SERVICE_URL). "
        "Ошибки доставки уведомления не блокируют сохранение сообщения (best-effort)."
    )
    shot(
        doc,
        "Рабочее веб-приложение: http://localhost:8002/ — форма входа/регистрации и список чатов (после логина).",
    )
    shot(doc, "Тот же UI: выбранный чат с отправленным сообщением.")
    shot(
        doc,
        "Опционально: Swagger chat-service — выполненный запрос POST .../messages и ответ 201 (или лог notification-service).",
    )

    add_heading(doc, "3.3. Документация микросервисов", 2)
    doc.add_paragraph(
        "Документация по архитектуре и ТЗ: docs/architecture.md, docs/technical_spec.md. "
        "Описание API каждого сервиса доступно интерактивно в Swagger UI (/docs) и в виде OpenAPI JSON. "
        "В корне репозитория — README.md и файл INSTRUKCIYA.md с инструкцией по запуску."
    )
    shot(doc, "Один из экранов /docs с раскрытой схемой запроса (Try it out) — на выбор.")

    doc.add_page_break()

    # ========== ЭТАП 5 ==========
    add_heading(doc, "4. Этап 5. Оптимизация, тестирование и развёртывание", 1)

    add_heading(doc, "4.1. Оптимизация", 2)
    doc.add_paragraph(
        "Выполнены меры, повышающие стабильность и предсказуемость запуска: разделение баз данных по "
        "сервисам; healthcheck PostgreSQL в Docker Compose; порядок старта зависимостей (ожидание готовности БД). "
        "Явное кеширование (Redis) и пагинация на уровне БД для тяжёлых выборок в объёме курсового проекта "
        "не внедрялись; сообщения запрашиваются с параметрами limit/offset."
    )

    add_heading(doc, "4.2. Тестирование (интеграция)", 2)
    doc.add_paragraph(
        "Интеграция компонентов проверяется запуском полного стека через docker compose и ручными/ "
        "автоматическими сценариями: регистрация → получение JWT → операции в chat-service. "
        "Юнит-тесты по сервисам изолированы (SQLite in-memory или тестовая конфигурация). "
        "При необходимости для отчёта можно добавить отдельный e2e-скрипт против поднятого Compose."
    )
    shot(
        doc,
        "По желанию: скрин успешного прохождения сценария в UI (два пользователя, чат, сообщения) или лог интеграционного теста.",
    )

    add_heading(doc, "4.3. Развёртывание в облаке и CI/CD", 2)
    doc.add_paragraph(
        "В текущей версии репозитория автоматическое развёртывание в облаке (Heroku, AWS, GCP) и "
        "pipeline CI/CD не настроены. Для промышленной сдачи рекомендуется: контейнеры в реестре "
        "(GHCR/ECR), оркестратор или PaaS, секреты вне образа, миграции Alembic, GitHub Actions — "
        "lint, pytest, build, deploy."
    )
    shot(
        doc,
        "Если будет настроен CI/CD или облако — вставить скриншот pipeline (успешный run) или панели облачного сервиса.",
    )

    add_heading(doc, "5. Заключение", 1)
    doc.add_paragraph(
        "Спроектирована и реализована микросервисная система чата с JWT, CRUD по чатам и сообщениям, "
        "сервисом уведомлений и контейнеризацией. Достигнуты ожидаемые результаты этапов по API, Docker и "
        "микросервисам; этап облака и CI/CD задокументирован как направление дальнейшего развития."
    )

    add_heading(doc, "Приложение А. Сводка: какие скриншоты вставить в отчёт", 1)
    bullets = [
        "Диаграмма архитектуры (из Mermaid или draw.io).",
        "Swagger user-service, chat-service, notification-service (по одному или композитно).",
        "Результаты pytest (один или несколько сервисов).",
        "Docker: контейнеры + фрагмент лога успешного старта.",
        "Веб-UI: главная после логина, чат с сообщениями.",
        "Опционально: Try it out в Swagger; облако/CI/CD — если появятся.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Number")

    doc.save(OUT_RU)
    doc.save(OUT_EN)
    print(f"Сохранено: {OUT_RU}")
    print(f"Сохранено: {OUT_EN}")


if __name__ == "__main__":
    main()
